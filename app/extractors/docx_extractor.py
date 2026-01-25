"""DOCX (Word document) content extraction."""

from typing import Dict, Any
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import io


class DOCXExtractor:
    """Extract text and metadata from DOCX files."""
    
    @staticmethod
    def extract_from_path(file_path: str, url: str) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            url: Source URL
            
        Returns:
            Dictionary with extracted data
        """
        try:
            doc = Document(file_path)
            return DOCXExtractor._extract_from_document(doc, url, file_path)
        except Exception as e:
            return {
                'title': 'Error extracting DOCX',
                'content': '',
                'metadata': {'error': str(e), 'url': url, 'file_path': file_path},
                'content_type': 'docx'
            }
    
    @staticmethod
    def extract_from_bytes(docx_bytes: bytes, url: str, filename: str = 'document.docx') -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX bytes.
        
        Args:
            docx_bytes: DOCX file content as bytes
            url: Source URL
            filename: Original filename
            
        Returns:
            Dictionary with extracted data
        """
        try:
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)
            return DOCXExtractor._extract_from_document(doc, url, filename)
        except Exception as e:
            return {
                'title': 'Error extracting DOCX',
                'content': '',
                'metadata': {'error': str(e), 'url': url, 'filename': filename},
                'content_type': 'docx'
            }
    
    @staticmethod
    def _extract_from_document(doc: Document, url: str, filename: str) -> Dict[str, Any]:
        """Extract from python-docx Document object."""
        text_content = []
        tables_data = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract from tables
        for table_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_rows.append(row_data)
            tables_data.append({
                'table_num': table_idx + 1,
                'rows': table_rows
            })
        
        # Get metadata
        metadata = DOCXExtractor._extract_metadata(doc)
        metadata['url'] = url
        metadata['paragraph_count'] = len(doc.paragraphs)
        metadata['table_count'] = len(doc.tables)
        
        # Extract title from document properties or filename
        title = metadata.get('title', '')
        if not title:
            title = filename.replace('.docx', '').replace('_', ' ').replace('-', ' ')
        
        return {
            'title': title or 'Untitled Document',
            'content': '\n\n'.join(text_content),
            'metadata': metadata,
            'content_type': 'docx',
            'tables': tables_data
        }
    
    @staticmethod
    def _extract_metadata(doc: Document) -> Dict[str, Any]:
        """Extract metadata from DOCX properties."""
        metadata = {}
        
        if doc.core_properties:
            props = doc.core_properties
            if props.title:
                metadata['title'] = props.title
            if props.author:
                metadata['author'] = props.author
            if props.subject:
                metadata['subject'] = props.subject
            if props.keywords:
                metadata['keywords'] = props.keywords
            if props.category:
                metadata['category'] = props.category
            if props.comments:
                metadata['comments'] = props.comments
            if props.created:
                metadata['created'] = props.created.isoformat()
            if props.modified:
                metadata['modified'] = props.modified.isoformat()
        
        return metadata
